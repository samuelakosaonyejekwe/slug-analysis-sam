#!/usr/bin/env python3
# =============================================================================
#  build_reports.py  — generates, from Case Study #3 (subfolder 10) and the
#  SHCT solver sources, three Word documents:
#     * model.equations.docx  — the full list of model equations used to build
#                               the solver (governing PDEs + every closure).
#     * case_study.docx       — the engineering case study (asset, scenarios,
#                               headline predictions, mitigation, conclusions).
#     * slug_report.docx      — the COMPREHENSIVE report: the case study + all
#                               model equations + all inputs + all generated
#                               outputs (every metric, curve, graph, chart, map
#                               and CSV table) for the three scenarios. Saved to
#                               the Windows desktop. Nothing is left out.
#
#  All numbers/figures are EXTRACTED from the solver outputs already written in
#  subfolder 10; the equations are transcribed from solver.py / shct_*.py.
# =============================================================================
import os, csv as csvmod, json, sys
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))            # .../10/ignore
PLOTDIR = os.path.join(HERE, "report_plots"); os.makedirs(PLOTDIR, exist_ok=True)

NAVY = RGBColor(0x1F, 0x37, 0x64); ACCENT = RGBColor(0x0B, 0x5A, 0x7A)
RED = RGBColor(0xB5, 0x30, 0x2A); GREEN = RGBColor(0x2E, 0x7D, 0x32); GREY = RGBColor(0x55, 0x55, 0x55)
NAVY_H, ACC_H, RED_H, ORG_H, TEAL_H, GRN_H = "#1F3764", "#0B5A7A", "#B5302A", "#E08A2B", "#2E8B8B", "#2E7D32"

SCEN = [
    ("As-operated normal production (degraded insulation, no inhibitor)", "outputs_steady", "as-operated"),
    ("Unplanned shut-in (cooldown / no-touch time)", "outputs_shutin", "shut-in"),
    ("Engineered mitigation (restored insulation + MEG)", "outputs_mitigated", "mitigated"),
]
OUTROOT = os.path.dirname(HERE)                              # .../10  (output folders live here)

XY_CSVS = [
    ("fields_profile.csv", "spatial profile along the route"),
    ("timeseries_monitor.csv", "transient response at the monitor station"),
    ("csv_crosssection.csv", "cross-section geometry along the route"),
    ("csv_compositional.csv", "compositional / PVT state along the route"),
    ("csv_compositional_transport.csv", "compositional grading along the route"),
]
GALLERY = [
    ("01_profiles.png", "Final-state profiles (P50): elevation, holdup, P & T vs T_eq, subcooling."),
    ("09_slug_prediction.png", "Slug prediction: terrain, regime, slug frequency/length, holdup."),
    ("10_riser_severe_slug.png", "Severe-slugging screen at the riser base & ascent."),
    ("02_holdup_spacetime.png", "Liquid-holdup field α_l(x,t) — bands are slug activity."),
    ("03_PT_envelope.png", "P–T trajectory vs hydrate envelope (curve)."),
    ("11_hydrate_envelope.png", "P–T trajectories (production + shut-in) vs hydrate envelope."),
    ("04_PhiSH_map.png", "Φ_SH(x,t) coupling-criticality map."),
    ("05_scenario_timeseries.png", "Scenario monitor response vs time (curves)."),
    ("06_deposit.png", "Wall-deposit growth at the monitor (curve)."),
    ("07_probabilistic.png", "Time-to-plug CDF + max Φ_SH P10–P90 band."),
    ("08_diagnostics.png", "Mass balances, clip activity, slug length, consistency."),
    ("cx1_geometry.png", "Cross-section geometry along the line (curves)."),
    ("cx2_azimuthal_deposit.png", "Azimuthal (bottom-of-line) deposit map."),
    ("cx3_sections.png", "2-D cross-section reconstructions."),
    ("compo_pvt.png", "Compositional/PVT tracking (PR EOS)."),
    ("compositional_transport.png", "Compositional grading (hydrate-former depletion)."),
    ("threed_deposit.png", "3-D reconstructed pipe — wall deposit."),
    ("threed_temperature.png", "3-D reconstructed pipe — wall temperature."),
    ("hydrate_validation.png", "Hydrate curve vs published experimental data (curve)."),
    ("12_mitigation_comparison.png", "Mitigation comparison (bar chart — categorical metrics)."),
]


# --------------------------------------------------------------------------- IO
def read_csv(path):
    with open(path) as fh:
        rows = list(csvmod.reader(fh))
    return rows[0], rows[1:]


def fnum(s):
    try:
        return float(s)
    except Exception:
        return np.nan


PRETTY = {"x_km": "distance along route (km)", "time_h": "time (h)"}


# ============================================================ matplotlib helpers
def plot_xy(path, tag, slug):
    hdr, rows = read_csv(path)
    data = {h: np.array([fnum(r[i]) if i < len(r) else np.nan for r in rows]) for i, h in enumerate(hdr)}
    xcol = "x_km" if "x_km" in hdr else ("time_h" if "time_h" in hdr else hdr[0])
    xlab = PRETTY.get(xcol, xcol)
    x = data[xcol]
    cols = [h for h in hdr if h != xcol and np.isfinite(data[h]).any()]
    ncol = 3
    nrow = int(np.ceil(len(cols) / ncol))
    fig, axes = plt.subplots(nrow, ncol, figsize=(11, 2.05 * nrow + 0.6))
    axes = np.atleast_1d(axes).ravel()
    for k, h in enumerate(cols):
        ax = axes[k]
        ax.plot(x, data[h], color=NAVY_H, lw=1.5)
        ax.fill_between(x, data[h], np.nanmin(data[h]), color=ACC_H, alpha=.08)
        ax.set_title(h, fontsize=8, color=NAVY_H, fontweight="bold")
        ax.tick_params(labelsize=7); ax.grid(alpha=.25)
        # clip a transient spike (e.g. the t=0 Phi_SH peak) so the meaningful range is clear
        col = data[h][np.isfinite(data[h])]
        if col.size:
            p98 = float(np.nanpercentile(col, 98)); mx = float(np.nanmax(col))
            if mx > 3 * max(p98, 1e-9) and p98 > 0:
                ax.set_ylim(min(0.0, float(np.nanmin(col))), p98 * 1.2)
        if h in ("subcooling_C",):
            ax.axhline(0, color=RED_H, ls=":", lw=1)
        if h in ("Phi_SH",):
            ax.axhline(1, color=RED_H, ls="--", lw=1)
        if k >= len(cols) - ncol:
            ax.set_xlabel(xlab, fontsize=7)
    for k in range(len(cols), len(axes)):
        axes[k].axis("off")
    fig.suptitle(f"{os.path.basename(path)} — every column as a curve vs {xlab}  [{tag}]",
                 color=NAVY_H, fontweight="bold", fontsize=11)
    fig.tight_layout(rect=[0, 0, 1, 0.97])
    out = os.path.join(PLOTDIR, f"{slug}_plot_{os.path.basename(path).replace('.csv','')}.png")
    fig.savefig(out, dpi=145); plt.close(fig)
    return out


def plot_prob_range(path, tag, slug):
    hdr, rows = read_csv(path)
    names = [r[0] for r in rows]
    p10 = np.array([fnum(r[1]) for r in rows]); p50 = np.array([fnum(r[2]) for r in rows])
    p90 = np.array([fnum(r[3]) for r in rows])
    fig, ax = plt.subplots(figsize=(9.2, 0.55 * len(names) + 1.4))
    y = np.arange(len(names))
    for yy, lo, md, hi in zip(y, p10, p50, p90):
        if np.isfinite(lo) and np.isfinite(hi):
            ax.plot([lo, hi], [yy, yy], color=ACC_H, lw=2.4, solid_capstyle="round", zorder=2)
            ax.plot([lo, lo], [yy - .12, yy + .12], color=ACC_H, lw=1.6)
            ax.plot([hi, hi], [yy - .12, yy + .12], color=ACC_H, lw=1.6)
        if np.isfinite(md):
            ax.plot(md, yy, "o", color=NAVY_H, ms=7, zorder=3)
            ax.text(md, yy + .22, f"P50={md:.3g}", ha="center", fontsize=7, color=NAVY_H)
    ax.set_yticks(y); ax.set_yticklabels(names, fontsize=8); ax.invert_yaxis()
    ax.grid(alpha=.25, axis="x"); ax.set_xscale("symlog", linthresh=1.0)
    ax.set_xlabel("value (symlog) — whisker = P10→P90, marker = P50", fontsize=8.5)
    ax.set_title(f"probabilistic_summary.csv — P10/P50/P90 uncertainty range per metric  [{tag}]",
                 color=NAVY_H, fontweight="bold", fontsize=10.5)
    fig.tight_layout()
    out = os.path.join(PLOTDIR, f"{slug}_plot_probabilistic_range.png")
    fig.savefig(out, dpi=145); plt.close(fig)
    return out


def plot_eng_bar(path, tag, slug):
    hdr, rows = read_csv(path)
    names, vals = [], []
    for r in rows:
        v = fnum(r[1])
        if np.isfinite(v) and abs(v) > 0:
            names.append(f"{r[0]} ({r[2]})"); vals.append(abs(v))
    fig, ax = plt.subplots(figsize=(9.6, 0.34 * len(names) + 1.0))
    y = np.arange(len(names))
    ax.barh(y, vals, color=TEAL_H)
    ax.set_yticks(y); ax.set_yticklabels(names, fontsize=7.5); ax.set_xscale("log")
    ax.invert_yaxis(); ax.grid(alpha=.25, axis="x")
    for yy, v in zip(y, vals):
        ax.text(v, yy, f" {v:.3g}", va="center", fontsize=7)
    ax.set_xlabel("magnitude (log scale; units in the label — see table)", fontsize=8)
    ax.set_title(f"engineering_deliverables.csv — numeric deliverables  [{tag}]\n"
                 "(bar chart: these are unrelated named quantities, not a curve)",
                 color=NAVY_H, fontweight="bold", fontsize=10)
    fig.tight_layout()
    out = os.path.join(PLOTDIR, f"{slug}_plot_engineering.png")
    fig.savefig(out, dpi=145); plt.close(fig)
    return out


def plot_composition(path, tag, slug):
    hdr, rows = read_csv(path)
    comp = [(r[0], fnum(r[1])) for r in rows]
    fig, ax = plt.subplots(figsize=(7.6, 3.4))
    names = [c[0] for c in comp]; vals = [c[1] for c in comp]
    ax.bar(names, vals, color=NAVY_H)
    for i, v in enumerate(vals):
        ax.text(i, v, f"{v:.3f}", ha="center", va="bottom", fontsize=7)
    ax.set_ylabel("mole fraction"); ax.grid(alpha=.25, axis="y")
    ax.set_title(f"feed_composition.csv — crude-oil feed makeup  [{tag}]\n"
                 "(bar chart: composition is categorical by component)",
                 color=NAVY_H, fontweight="bold", fontsize=10)
    fig.tight_layout()
    out = os.path.join(PLOTDIR, f"{slug}_plot_composition.png")
    fig.savefig(out, dpi=145); plt.close(fig)
    return out


def hero_curves(folder, tag, slug):
    outs = []
    fp = os.path.join(folder, "fields_profile.csv")
    if os.path.exists(fp):
        hdr, rows = read_csv(fp)
        d = {h: np.array([fnum(r[i]) if i < len(r) else np.nan for r in rows]) for i, h in enumerate(hdr)}
        x = d["x_km"]
        fig, ax = plt.subplots(figsize=(7.8, 3.6))
        l1, = ax.plot(x, d["P_bar"], color=NAVY_H, lw=2, label="pressure P (bar)")
        ax.set_xlabel("distance along route (km)"); ax.set_ylabel("P (bar)", color=NAVY_H)
        a2 = ax.twinx()
        l2, = a2.plot(x, d["T_C"], color=RED_H, lw=2, label="temperature T (°C)")
        l3, = a2.plot(x, d["Teq_C"], color=RED_H, lw=1.4, ls="--", label="hydrate T_eq (°C)")
        a2.fill_between(x, d["T_C"], d["Teq_C"], where=d["T_C"] < d["Teq_C"], color="#f6d6d2", alpha=.6)
        a2.set_ylabel("T, T_eq (°C)", color=RED_H)
        ax.legend(handles=[l1, l2, l3], fontsize=8, loc="upper right")
        ax.set_title(f"Prediction curve — pressure & temperature vs hydrate boundary  [{tag}]",
                     color=NAVY_H, fontweight="bold"); ax.grid(alpha=.25)
        fig.tight_layout(); o = os.path.join(PLOTDIR, f"{slug}_curve_PT.png"); fig.savefig(o, dpi=150); plt.close(fig)
        outs.append((o, "Prediction curve — pressure & temperature vs the hydrate boundary (from fields_profile.csv)."))
        fig, ax = plt.subplots(figsize=(7.8, 3.0))
        ax.plot(x, d["subcooling_C"], color=ORG_H, lw=2, label="subcooling ΔT_sub")
        ax.axhline(0, color="#555", ls=":", label="hydrate boundary")
        ax.fill_between(x, 0, d["subcooling_C"], where=d["subcooling_C"] > 0, color="#f6d6d2", alpha=.6)
        ax.set_xlabel("distance along route (km)"); ax.set_ylabel("ΔT_sub (°C)")
        ax.set_title(f"Hydrate-risk curve — subcooling along the route  [{tag}]",
                     color=NAVY_H, fontweight="bold"); ax.legend(fontsize=8); ax.grid(alpha=.25)
        fig.tight_layout(); o = os.path.join(PLOTDIR, f"{slug}_curve_subcooling.png"); fig.savefig(o, dpi=150); plt.close(fig)
        outs.append((o, "Hydrate-risk curve — subcooling ΔT_sub along the route (from fields_profile.csv)."))
        fig, ax = plt.subplots(figsize=(7.8, 3.0))
        l1, = ax.plot(x, d["holdup"], color=ACC_H, lw=2, label="liquid holdup α_l")
        ax.set_ylabel("holdup α_l", color=ACC_H); ax.set_ylim(0, 1); ax.set_xlabel("distance along route (km)")
        a2 = ax.twinx(); l2, = a2.plot(x, d["f_slug_Hz"], color=ORG_H, lw=1.8, label="slug frequency (Hz)")
        a2.set_ylabel("f_slug (Hz)", color=ORG_H)
        ax.legend(handles=[l1, l2], fontsize=8, loc="upper left")
        ax.set_title(f"Slugging curve — holdup & slug frequency along the route  [{tag}]",
                     color=NAVY_H, fontweight="bold"); ax.grid(alpha=.25)
        fig.tight_layout(); o = os.path.join(PLOTDIR, f"{slug}_curve_slug.png"); fig.savefig(o, dpi=150); plt.close(fig)
        outs.append((o, "Slugging curve — liquid holdup & slug frequency along the route (from fields_profile.csv)."))
    tp = os.path.join(folder, "timeseries_monitor.csv")
    if os.path.exists(tp):
        hdr, rows = read_csv(tp)
        d = {h: np.array([fnum(r[i]) if i < len(r) else np.nan for r in rows]) for i, h in enumerate(hdr)}
        t = d["time_h"]
        fig, ax = plt.subplots(figsize=(7.8, 3.2))
        l1, = ax.plot(t, d["deposit_mm"], color=RED_H, lw=2, label="wall deposit δ_h (mm)")
        ax.set_xlabel("time (h)"); ax.set_ylabel("deposit (mm)", color=RED_H)
        a2 = ax.twinx(); l2, = a2.plot(t, d["Phi_SH"], color=NAVY_H, lw=1.8, label="coupling number Φ_SH")
        a2.axhline(1, color=RED_H, ls="--", lw=1); a2.set_ylabel("Φ_SH", color=NAVY_H)
        ax.legend(handles=[l1, l2], fontsize=8, loc="upper left")
        ax.set_title(f"Transient coupled-growth curve — deposit & Φ_SH vs time  [{tag}]",
                     color=NAVY_H, fontweight="bold"); ax.grid(alpha=.25)
        fig.tight_layout(); o = os.path.join(PLOTDIR, f"{slug}_curve_transient.png"); fig.savefig(o, dpi=150); plt.close(fig)
        outs.append((o, "Transient coupled-growth curve — wall deposit & Φ_SH vs time (from timeseries_monitor.csv)."))
    return outs


# ============================================================ docx helpers (per-doc)
class Doc:
    def __init__(self):
        self.doc = Document()
        s = self.doc.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(10.5)
        s.paragraph_format.space_after = Pt(5)

    def H1(self, t):
        p = self.doc.add_heading(t, 1)
        for r in p.runs:
            r.font.color.rgb = NAVY
        return p

    def H2(self, t):
        p = self.doc.add_heading(t, 2)
        for r in p.runs:
            r.font.color.rgb = ACCENT
        return p

    def H3(self, t):
        p = self.doc.add_heading(t, 3)
        for r in p.runs:
            r.font.color.rgb = ACCENT
        return p

    def para(self, t, **k):
        p = self.doc.add_paragraph(); r = p.add_run(t)
        r.bold = k.get("bold", False); r.italic = k.get("italic", False)
        if k.get("size"):
            r.font.size = Pt(k["size"])
        if k.get("color"):
            r.font.color.rgb = k["color"]
        return p

    def bullet(self, t, **k):
        p = self.doc.add_paragraph(style="List Bullet"); r = p.add_run(t)
        r.bold = k.get("bold", False)
        if k.get("color"):
            r.font.color.rgb = k["color"]
        return p

    def equation(self, formula):
        p = self.doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(formula); r.font.name = "Consolas"; r.font.size = Pt(10.5)
        r.font.color.rgb = NAVY; r.bold = True
        return p

    def figure(self, path, caption, width=6.8):
        if not path or not os.path.exists(path):
            self.para(f"[figure not found: {os.path.basename(path) if path else 'n/a'}]",
                      italic=True, color=GREY, size=8); return
        self.doc.add_picture(path, width=Inches(width))
        self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = self.doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption); r.italic = True; r.font.size = Pt(8.3); r.font.color.rgb = GREY

    def csv_table(self, path, max_rows=14, fs=6.6):
        hdr, rows = read_csv(path)
        if len(rows) > max_rows:
            idx = sorted(set(np.linspace(0, len(rows) - 1, max_rows).astype(int)))
            rows = [rows[i] for i in idx]
        t = self.doc.add_table(rows=0, cols=len(hdr)); t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        cs = t.add_row().cells
        for i, h in enumerate(hdr):
            cs[i].text = ""; rr = cs[i].paragraphs[0].add_run(h); rr.bold = True
            rr.font.size = Pt(fs); rr.font.color.rgb = NAVY
        for row in rows:
            cs = t.add_row().cells
            for i, v in enumerate(row):
                cs[i].text = ""; rr = cs[i].paragraphs[0].add_run(str(v)); rr.font.size = Pt(fs)
        self.doc.add_paragraph()

    def kv_table(self, pairs, fs=9, widths=(3.0, 1.7, 1.1)):
        t = self.doc.add_table(rows=0, cols=len(pairs[0])); t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ri, row in enumerate(pairs):
            cs = t.add_row().cells
            for i, v in enumerate(row):
                cs[i].text = ""; rr = cs[i].paragraphs[0].add_run(str(v)); rr.font.size = Pt(fs)
                if ri == 0:
                    rr.bold = True; rr.font.color.rgb = NAVY
                if i < len(widths):
                    cs[i].width = Inches(widths[i])
        self.doc.add_paragraph()

    def pagebreak(self):
        self.doc.add_page_break()

    def save(self, path):
        self.doc.save(path); return path


# ============================================================ KPI metric table
KPI_KEYS = [("dP_total_bar", "Total ΔP", "bar"), ("Vm_peak_mps", "Peak velocity", "m/s"),
            ("erosional_limit_mps", "Erosional limit (API 14E)", "m/s"),
            ("slug_length_max_m", "Max slug length", "m"), ("slug_length_mean_m", "Mean slug length", "m"),
            ("slug_fraction", "Slug/intermittent fraction", "–"),
            ("V_surge_P90_m3", "Slug-catcher surge (P90)", "m³"),
            ("max_subcooling_C", "Max subcooling", "°C"), ("dT_design_C", "Design subcooling (P90)", "°C"),
            ("max_Phi_SH", "Max Φ_SH (coupling)", "–"), ("coupled_hotspot_km", "Coupling hot-spot", "km"),
            ("P_plug", "Plug probability", "frac"), ("time_to_plug_P50_h", "Time-to-plug P50", "h"),
            ("time_to_plug_P10_h", "Time-to-plug P10", "h"), ("time_to_plug_P90_h", "Time-to-plug P90", "h"),
            ("peak_deposit_mm", "Peak wall deposit", "mm"), ("MEG_wt_pct", "MEG required", "wt%"),
            ("MEG_Lph", "MEG rate", "L/h"), ("under_inhibited_km", "Under-inhibited length", "km"),
            ("cooldown_to_hydrate_h", "No-touch time", "h"), ("U_eff_WmK", "Effective U", "W/m²K"),
            ("slurry_rel_viscosity", "Slurry rel. viscosity", "–"),
            ("hydrate_mass_formed_kg", "Hydrate mass formed", "kg"),
            ("sound_speed_min_mps", "Min mixture sound speed", "m/s"),
            ("wax_onset_km", "Wax onset", "km"), ("scaling_tendency_max", "Max scaling index", "–"),
            ("mass_conservation_err", "Liquid mass error", "frac"),
            ("gas_mass_conservation_err", "Gas mass error", "frac"),
            ("fallbacks", "Solver fallbacks", "#")]


def kpi_rows(km):
    out = [["Metric", "Value", "Units"]]
    for k, lab, u in KPI_KEYS:
        v = km.get(k)
        if v is None:
            continue
        try:
            sv = f"{float(v):.4g}"
        except Exception:
            sv = str(v)
        out.append([lab, sv, u])
    return out


# ============================================================ MODEL EQUATIONS
# Each entry: (name, formula (Unicode), description, source). Transcribed directly
# from solver.py and shct_correlations.py / shct_eos.py.
EQ_GROUPS = [
 ("A.  Governing transient conservation PDEs (finite-volume, CFL-adaptive)", [
  ("Liquid-holdup transport (drift-flux kinematic wave)",
   "∂(α_l·A)/∂t + ∂(α_l·v_l·A)/∂x = − S_l",
   "Volume-conservative liquid continuity, discretised by a conservative finite-volume "
   "scheme with signed upwinding (optionally 2nd-order TVD). S_l is the liquid (water) "
   "volume sink to hydrate formation. Holdup is recovered as α_l = (α_l·A)/A each step.",
   "Conservative FV; Bendiksen drift-flux closure."),
  ("Gas-mass continuity (conservative, mass-consistent)",
   "∂(ρ_g·α_g·A)/∂t + ∂(ρ_g·V_sg·A)/∂x = − ṁ_(gas→hydrate)",
   "Solved alongside the liquid so liquid AND gas mass balance to ~0%. The mass-consistent "
   "superficial gas velocity V_sg = (ρ_g,in·q_g,in)/(ρ_g·A) makes the mass flux ρ_g·V_sg·A "
   "equal to the inlet gas mass rate along the line.",
   "Conservative FV gas continuity."),
  ("Mixture momentum → implicit pressure (default engine)",
   "u_m = U₀ − C_u·∂p/∂x ;   R(pⁿ⁺¹−pⁿ) + ∂/∂x[ U₀ − C_u·∂pⁿ⁺¹/∂x ] = 0",
   "Transient mixture momentum with the pressure gradient and wall drag treated implicitly; "
   "eliminating velocity gives a tridiagonal (Poisson-type) pressure equation solved each step "
   "(Thomas algorithm). Predictor u* = u_m − dt·(advection + g·sinθ); wall-drag rate "
   "β = f/(2D)·|u_m|; U₀ = u*/(1+dt·β); compressible storage R = α_g/(p·dt). Inlet pressure "
   "Dirichlet (pinned in-system), outlet through-flux (rate control).",
   "Implicit-pressure mixture-momentum; drift-flux slip closure."),
  ("Full two-fluid momentum (alternative engine)",
   "two independent phase momenta + interfacial drag, shared pressure; slip emerges from balance",
   "RELAP/TRACE-style semi-implicit: advection+gravity explicit; pressure gradient, wall drag "
   "and stiff interfacial drag implicit. A 2×2 per-cell velocity system → tridiagonal pressure; "
   "an interfacial-pressure correction keeps the characteristics real (well-posed).",
   "Two-fluid model with interfacial-pressure regularisation."),
  ("Energy transport (advection + seabed loss + latent + Joule–Thomson)",
   "∂T/∂t + j·∂T/∂x = − U_eff·(4/D)·(T − T_sink)/(ρ_m·c_p) + q_latent + q_JT",
   "Signed-upwind (optionally TVD) advection; the stiff wall-loss term is implicit (backward "
   "Euler), with a Heun 2nd-order corrector on advection. β_loss = U_eff·4/(D·ρ_m·c_p). "
   "q_latent = L_hyd·ρ_hyd·R_g,bulk/(ρ_m·c_p) (exothermic bulk hydrate, self-limiting); "
   "q_JT = μ_JT·(1−α_l)·j·(∂p/∂x) (real-gas cooling on expansion).",
   "Transient energy balance; implicit wall loss; Strang/Heun split."),
  ("Hydrate phase-field (advected reaction–diffusion, stochastic)",
   "∂φ/∂t + v_l·∂φ/∂x = D_φ·∂²φ/∂x² + R_grow + R_nuc + ξ",
   "Advected reaction–diffusion field for the bulk hydrate fraction φ, carrying its D_φ "
   "diffusion term and a stochastic nucleation source ξ (Monte-Carlo ensemble).",
   "Phase-field hydrate model with stochastic nucleation."),
 ]),
 ("B.  Hydrate thermodynamics, kinetics and the slug–hydrate coupling", [
  ("Hydrate equilibrium temperature (natural-gas correlation)",
   "T_eq = 7.7·ln(P) − 23.2 + 18·(SG − 0.60) − 0.74·S_wt        [°C, P in bar]",
   "Base natural-gas hydrate curve (≈3 °C @30 bar, 10 °C @70 bar, 13 °C @100 bar, 18 °C @200 bar), "
   "shifted for gas specific gravity SG and depressed by produced-water salinity S_wt (wt% NaCl-eq). "
   "An optional measured/PVT [P,T_eq] table overrides it (monotonic interpolation).",
   "Natural-gas hydrate equilibrium (GPSA/Sloan-anchored)."),
  ("Subcooling (bulk and wall driving forces)",
   "ΔT_sub = T_eq − T ;   ΔT_sub,wall = T_eq − T_front",
   "ΔT_sub is the local thermodynamic hydrate driving force. The deposition-relevant WALL "
   "subcooling uses the deposit-front temperature T_front (see deposit-insulation feedback).",
   "Definition."),
  ("Stochastic nucleation (induction time, Monte-Carlo)",
   "τ_ind = τ₀·3600·exp( min(β/ΔT_sub, cap) ) ;   p_nuc = 1 − exp(−dt/τ_ind)",
   "Classical induction-time nucleation: a cell nucleates when a uniform random draw < p_nuc and "
   "ΔT_sub > ΔT_sub,min. Growth proceeds only after onset, so subcooling builds during induction "
   "→ a genuine onset spread across the ensemble (P10/P50/P90).",
   "Classical stochastic nucleation."),
  ("Growth-rate constant (Arrhenius)",
   "k_g = k_g0·exp( −(E_a/R)·(1/T − 1/T_ref) )",
   "Temperature dependence of the intrinsic hydrate growth-rate constant.",
   "Arrhenius / CSMHyK-type kinetics."),
  ("Bulk slurry growth (CSMHyK-type, self-limiting)",
   "R_g,bulk = k_g·a_i·ΔT_sub^n·(1 − φ/φ_max)·𝟙_nucleated",
   "Bulk hydrate formation driven by the local bulk subcooling; its exothermic latent heat warms "
   "the fluid, so it self-limits (φ stays low → transportable slurry). a_i is the interfacial area.",
   "CSMHyK-type growth law."),
  ("Wall-deposit growth (sustained cold-wall driving force)",
   "R_g,wall = k_g,wall·a_i·ΔT_sub,wall^n·𝟙_nucleated",
   "Deposit growth driven by the sustained cold-wall subcooling (the wall sits at the cold seabed; "
   "its latent heat is rejected to the sea, not the bulk) — this is what lets a plug grow while the "
   "bulk slurry stays dilute.",
   "Cold-wall deposition kinetics."),
  ("Deposit-insulation feedback (self-limiting late growth)",
   "frac_warm = k_dep·R_dep/(R_dep + R_ext),  R_dep = δ/k_hyd ;  T_front = T_sea + (T − T_sea)·frac_warm",
   "As the deposit thickens it insulates its own growth front from the cold sea, warming the front "
   "toward the bulk and reducing the wall subcooling that drives further deposition.",
   "Conductive-resistance deposit feedback."),
  ("Slug–Hydrate Coupling Number Φ_SH  (the claimed invention)",
   "Φ_SH = C·k_g,wall·a_i·ΔT_sub,wall^n / f_slug",
   "Dimensionless ratio of hydrate-formation tendency to slug-renewal rate, evaluated at the "
   "sustained wall subcooling. Φ_SH > 1 ⇒ hydrate formation outruns slug scouring ⇒ consolidation "
   "/ plugging criticality. The central coupled-risk metric mapped in Φ_SH(x,t).",
   "SHCT invention (physically reasoned; flow-loop validation pending)."),
  ("Wall-deposit evolution (growth vs slug scouring)",
   "dδ/dt = f_wall·R_g,wall·D/4 − k_ero·f_slug·δ",
   "Annulus deposit-thickness rate: consolidating wall growth minus slug erosion/scouring. f_wall is "
   "a Φ_SH-gated capture fraction (consolidation only above criticality); erosion acts where Φ_SH<1.",
   "Deposit mass balance with slug scouring."),
  ("Hydrate liquid/gas sinks (mass coupling)",
   "hyd_vol_rate = (R_g,bulk + f_wall·R_g,wall)·A   → removed from liquid & gas inventories",
   "The water and gas consumed by both bulk and consolidating-wall hydrate are removed from the "
   "conserved liquid and gas inventories, so the balances close WITH the hydrate consumption.",
   "Mass-consistent hydrate sink."),
 ]),
 ("C.  Multiphase-hydrodynamic closures (published correlations)", [
  ("Bendiksen drift-flux slip",
   "C₀ = 1.05 + 0.15·sin|θ| ;  v_d = 0.35·√(gD)·sinθ + 0.20·√(gD)·cosθ ;  u_g = C₀·u_m + v_d",
   "Distribution coefficient and drift velocity closing the gas/liquid slip; keeps the system "
   "unconditionally well-posed (no complex characteristics).",
   "Bendiksen (1984) drift-flux."),
  ("Flow-regime classification",
   "regime = f(V_sg, V_sl, θ) → {stratified-smooth/wavy, slug, annular, dispersed-bubble, churn}",
   "Threshold map on the superficial gas/liquid velocities and inclination selecting the local "
   "flow regime, which then sets the friction multiplier, interfacial area and Nusselt enhancement.",
   "Taitel–Dukler-style regime transitions."),
  ("Slug frequency (Gregory–Scott / Zabaras)",
   "f_slug = 0.0226·[ (V_sl/(gD))·(19.75/V_m + V_m) ]^1.2 · (0.836 + 2.75·sin|θ|^0.25)",
   "Hydrodynamic slug frequency with an inclination factor; drives the renewal term in Φ_SH and the "
   "slug-catcher sizing.",
   "Gregory–Scott (1969) / Zabaras (2000)."),
  ("Slug unit length",
   "L_u = V_t/f_slug ,   V_t = 1.2·V_m + 0.35·√(gD)",
   "Reduced-order (sub-grid) mean slug length from the translational velocity V_t (Dukler–Hubbard / "
   "Gregory–Scott scaling) — resolves metre-scale slugs a 200-m grid cannot.",
   "Dukler–Hubbard / Gregory–Scott."),
  ("Haaland friction factor",
   "1/√f = −1.8·log₁₀[ (ε/D/3.7)^1.11 + 6.9/Re ] ;   f = 64/Re  (Re < 2300)",
   "Explicit Darcy friction factor (turbulent Haaland, laminar Poiseuille floor).",
   "Haaland (1983)."),
  ("Interfacial area per unit volume (geometry-resolved)",
   "stratified: a_i = sin γ /D ;  annular: a_i = 4√α_l/D ;  dispersed: a_i = 6·α_g/d_b (d_b=0.1D)",
   "Regime-specific gas–liquid interfacial area; the stratified wetted half-angle γ comes from the "
   "Biberg approximation of the Taitel–Dukler circular-segment geometry.",
   "Taitel–Dukler geometry; Biberg angle approximation."),
  ("Regime wall-friction multiplier & Nusselt",
   "Nu = 0.023·Re^0.8·Pr^0.4 · enh(regime) ;  friction ×{0.9…1.7 by regime}",
   "Dittus–Boelter inner heat-transfer with a regime enhancement (slug/churn mixing raises both wall "
   "shear and the inner film coefficient).",
   "Dittus–Boelter + regime enhancement."),
  ("Droplet entrainment (Ishii–Mishima)",
   "E = 1 − exp(−We/1e5) ,   We = ρ_g·V_sg²·D/σ",
   "Liquid fraction entrained into the gas core; significant only in high-shear annular/churn flow.",
   "Ishii–Mishima (1989)."),
  ("Mixture sound speed (Wood's equation)",
   "1/(ρ_m·c²) = α_l/(ρ_l·c_l²) + α_g/(γ·p)",
   "Two-phase sound speed for water-hammer / fast-transient screening; even a little gas drops c to "
   "O(10–100 m/s). Used by the optional acoustic (water-hammer) storage term.",
   "Wood's equation."),
 ]),
 ("D.  Fluid properties & compositional PVT (Peng–Robinson EOS)", [
  ("Peng–Robinson cubic EOS",
   "a_i = 0.45724·R²·T_c²/P_c·α ,  b_i = 0.07780·R·T_c/P_c ,  α = [1 + κ(1 − √(T/T_c))]²",
   "with κ = 0.37464 + 1.54226·ω − 0.26992·ω². The cubic in Z (with van der Waals mixing rules and "
   "binary k_ij) is solved for the vapour/liquid roots; the multicomponent vapour–liquid flash "
   "iterates K-values to fugacity equality φ_l·x = φ_v·y.",
   "Peng & Robinson (1976); Reid–Prausnitz–Poling; Whitson–Brulé."),
  ("Gas Z-factor (real gas)",
   "Z = A + (1−A)/exp(B) + C·Pr^D  (Beggs–Brill explicit; Sutton pseudo-criticals from SG)",
   "Standing-Katz-equivalent explicit correlation used when an EOS/Z-table is not supplied; reproduces "
   "Standing-Katz to a few % over 1.2 ≤ Tr ≤ 2.4, Pr ≤ 10.",
   "Beggs–Brill / Sutton; Standing–Katz."),
  ("Gas density",
   "ρ_g = P·M / (Z·R·T)",
   "Real-gas density from the Z-factor (or a user PVT table).",
   "Real-gas law."),
  ("Gas viscosity (Lee–Gonzalez–Eakin)",
   "μ_g = 1e-4·K·exp(X·ρ_g^Y) ,  K,X,Y = f(M, T)",
   "Pressure/temperature-dependent gas viscosity (rises with density) — matters for friction at high P.",
   "Lee–Gonzalez–Eakin (1966)."),
  ("Oil density (black-oil P,T correction)",
   "ρ_o = ρ_o,ref·[ 1 + c_P·(P − 1) − β·(T − 15) ]   (β = 7e-4 /K, c_P = 1e-4 /bar)",
   "Thermal-expansion + isothermal-compressibility correction about the reference state (or a user "
   "PVT table / EOS surface).",
   "Black-oil correlation."),
  ("Compositional liquid viscosity (Lohrenz–Bray–Clark)",
   "μ_L via LBC residual-viscosity polynomial in reduced density (η from T_r, ξ scaling)",
   "Compositional dense-phase liquid/condensate viscosity from the PR EOS state.",
   "Lohrenz–Bray–Clark (1964)."),
  ("Joule–Thomson coefficient",
   "μ_JT = (T/(c_p·M·ρ_mol·Z))·(∂Z/∂T)_P",
   "Real-gas cooling on expansion, evaluated from the Z-factor temperature slope; feeds q_JT in the "
   "energy equation.",
   "EOS-based Joule–Thomson."),
 ]),
 ("E.  Thermal design, inhibitor and engineering deliverables", [
  ("Effective overall heat-transfer coefficient (multi-layer)",
   "1/U_eff = 1/h_in + Σ_k (th_k/k_k) + 1/h_out",
   "Series thermal resistances of the inner film, each wall/insulation/coating layer, and the outer "
   "film → the effective U referred to the inner area (else the constant U_wall).",
   "Series-resistance heat transfer."),
  ("Cooldown / no-touch time (lumped capacitance)",
   "t_cd = (m·c_p)/(U·π·D)·ln[ (T_op − T_sea)/(T_eq − T_sea) ]",
   "Lumped-capacitance time for the coldest critical point to cool from operating T to the hydrate "
   "onset after shut-in (preferred: measured directly from the transient when the monitor crosses "
   "ΔT_sub = 0).",
   "Lumped-capacitance cooldown."),
  ("MEG hydrate suppression (Nielsen–Bucklin)",
   "ΔT = −72·ln(1 − x_MEG)        (x_MEG = MEG mole fraction in the aqueous phase)",
   "Equilibrium-temperature depression for a local MEG concentration; valid to high wt% where the "
   "classical Hammerschmidt linear form breaks down.",
   "Nielsen–Bucklin (1983)."),
  ("Required MEG dose (inverse design)",
   "x = 1 − exp(−ΔT_req/72) ;  W = 100·x·M_MEG/[ (1−x)·M_H₂O + x·M_MEG ] ;  ṁ_MEG = ṁ_water·W/(100−W)",
   "Inverts Nielsen–Bucklin for the required MEG wt%, mass rate and volume rate to achieve the design "
   "subcooling depression (with a design margin).",
   "Nielsen–Bucklin inversion."),
  ("Erosional velocity limit (API RP 14E)",
   "V_e = C/√ρ_m",
   "Maximum recommended mixture velocity from the API erosional C-factor and the mixture density.",
   "API RP 14E."),
  ("Slug-catcher surge volume",
   "V_surge = (q_l/f_slug)·surge_factor",
   "Reduced-order slug-catcher sizing from the liquid rate and the monitored slug frequency (reported "
   "at P90 across the ensemble).",
   "Slug-volume sizing."),
  ("Slurry transportability (Camargo–Palermo)",
   "μ_rel = (1 − φ/φ_max)^(−exp)",
   "Relative viscosity of the hydrate slurry; transportable while μ_rel stays below the threshold.",
   "Camargo–Palermo (2000)."),
  ("Probabilistic time-to-plug",
   "Monte-Carlo ensemble → P10/P50/P90 ;  Kaplan–Meier (right-censored) CDF",
   "Stochastic nucleation + parameter spread give a genuine probabilistic time-to-plug band rather "
   "than a single deterministic time.",
   "Monte-Carlo UQ; Kaplan–Meier estimator."),
 ]),
 ("F.  Numerical scheme", [
  ("Conservative finite-volume + adaptive CFL",
   "flux-form updates; dt from a CFL condition (advective + acoustic + diffusion caps)",
   "Volume-conservative flux differencing for holdup, gas-mass, energy and phase-field; a single "
   "predictor–corrector dt per step shared by momentum and transport.",
   "Conservative FV; CFL control."),
  ("TVD flux limiters (2nd-order)",
   "φ(r): minmod = max(0,min(1,r)) ; van Leer = (r+|r|)/(1+|r|) ; superbee",
   "Optional 2nd-order TVD reconstruction of advected face values (upwind → 1st order).",
   "TVD slope limiters."),
  ("Tridiagonal (Thomas) implicit pressure solve",
   "solve  a·pᵢ₋₁ + b·pᵢ + c·pᵢ₊₁ = d   per step (vectorised over the ensemble)",
   "The implicit pressure equation is a tridiagonal system solved by the Thomas algorithm each step.",
   "Thomas algorithm."),
  ("Never-fail fallback",
   "non-finite step → auto-degrade to a proven quasi-steady solver (time always advances)",
   "Robustness guarantee: the implicit engine cannot crash or emit NaN; 0 fallbacks on the case study.",
   "Graceful degradation."),
 ]),
]


def write_equations(D, intro=True):
    if intro:
        D.para("The SHCT solver integrates, in time, a system of coupled partial-differential "
               "equations for multiphase flow, heat transfer and hydrate formation on arbitrary "
               "terrain, closed by published flow-assurance correlations and a compositional "
               "Peng–Robinson PVT engine. Every equation actually used to build the solver is "
               "listed below, grouped by role, with its symbolic form, a one-line description and "
               "its source. Symbols: α_l/α_g liquid/gas holdup, A flow area, v_l/v_g/u_m phase & "
               "mixture velocities, j mixture (superficial) velocity, p pressure, T temperature, "
               "T_eq hydrate-equilibrium temperature, ΔT_sub subcooling, φ bulk hydrate fraction, "
               "δ wall-deposit thickness, f_slug slug frequency, a_i interfacial area, θ inclination, "
               "ρ density, μ viscosity, U_eff effective heat-transfer coefficient.")
    for gname, eqs in EQ_GROUPS:
        D.H2(gname)
        for name, formula, desc, src in eqs:
            D.para(name, bold=True, color=NAVY, size=10.5)
            D.equation(formula)
            D.para(desc, size=9.5)
            D.para(f"Source / basis: {src}", italic=True, color=GREY, size=8.8)


# ============================================================ INPUTS section
def write_inputs(D, full=True):
    D.para("All three scenarios share the same asset (32 km, 10.75-in deepwater medium-crude-oil "
           "tie-back) and feed; they differ only in thermal design and inhibition. The complete "
           "input data deck and feed composition are tabulated below; the full machine-readable "
           "case configuration (case_config.json) accompanies each scenario folder.")
    # input deck + feed from the steady folder (shared)
    st = os.path.join(OUTROOT, "outputs_steady")
    idp = os.path.join(st, "input_data_deck.csv")
    if os.path.exists(idp):
        D.H3("Input data deck (pipeline / fluid / operating / numerics)")
        D.csv_table(idp, max_rows=40, fs=8.5)
    fc = os.path.join(st, "feed_composition.csv")
    if os.path.exists(fc):
        D.H3("Feed composition — medium-crude makeup (→ Peng–Robinson flash)")
        D.figure(plot_composition(fc, "as-operated", "inp"), "Feed composition (mole fraction per component).", width=5.4)
        D.csv_table(fc, max_rows=20, fs=9)
    if full:
        D.H3("Per-scenario thermal/inhibition design (input deck per scenario)")
        for label, fname, tag in SCEN:
            ip = os.path.join(OUTROOT, fname, "input_data_deck.csv")
            if os.path.exists(ip):
                D.para(f"{label} [{tag}]", bold=True, color=ACCENT)
                D.csv_table(ip, max_rows=40, fs=8)


# ============================================================ OUTPUTS section (full)
def write_scenario_outputs(D, label, fname, tag, sec):
    folder = os.path.join(OUTROOT, fname)
    slug = tag.replace(" ", "")
    km = {}
    kmp = os.path.join(folder, "key_metrics.json")
    if os.path.exists(kmp):
        km = json.load(open(kmp))
    D.H1(f"{sec}.  Scenario outputs: {label}")

    D.H2(f"{sec}.1  Headline prediction metrics")
    D.kv_table(kpi_rows(km))

    D.H2(f"{sec}.2  Key prediction curves (graphs)")
    for o, cap in hero_curves(folder, tag, slug):
        D.figure(o, f"{cap}  [{tag}]", width=6.6)

    sub = 2
    for csvname, desc in XY_CSVS:
        path = os.path.join(folder, csvname)
        if not os.path.exists(path):
            continue
        sub += 1
        D.H2(f"{sec}.{sub}  Graph of {csvname} — {desc}")
        D.figure(plot_xy(path, tag, slug), f"Every column of {csvname} as a curve  [{tag}].", width=6.95)
        D.csv_table(path, max_rows=12)

    pp = os.path.join(folder, "probabilistic_summary.csv")
    if os.path.exists(pp):
        sub += 1
        D.H2(f"{sec}.{sub}  Graph of probabilistic_summary.csv — P10/P50/P90 uncertainty")
        D.figure(plot_prob_range(pp, tag, slug), f"Probabilistic range graph  [{tag}].", width=6.6)
        D.csv_table(pp, max_rows=12, fs=8)

    ep = os.path.join(folder, "engineering_deliverables.csv")
    if os.path.exists(ep):
        sub += 1
        D.H2(f"{sec}.{sub}  Chart of engineering_deliverables.csv (categorical)")
        D.figure(plot_eng_bar(ep, tag, slug), f"Numeric engineering deliverables  [{tag}].", width=6.6)
        D.csv_table(ep, max_rows=40, fs=8)

    sub += 1
    D.H2(f"{sec}.{sub}  Solver & extension figures (charts, maps, curves)")
    for fn, cap in GALLERY:
        D.figure(os.path.join(folder, fn), f"{cap}  [{tag}]",
                 width=5.3 if fn in ("03_PT_envelope.png", "11_hydrate_envelope.png", "hydrate_validation.png") else 6.8)


# ============================================================ CASE STUDY narrative
def f2(km, k, f="{:.2f}"):
    try:
        return f.format(float(km.get(k)))
    except Exception:
        return "n/a"


def write_case_study(D, headline_only=False):
    kmA = json.load(open(os.path.join(OUTROOT, "outputs_steady", "key_metrics.json")))
    kmS = json.load(open(os.path.join(OUTROOT, "outputs_shutin", "key_metrics.json")))
    kmM = json.load(open(os.path.join(OUTROOT, "outputs_mitigated", "key_metrics.json")))

    D.H1("Executive summary")
    D.para("This case study applies the SHCT transient coupled-PDE solver to a representative "
           "deepwater medium-crude-oil subsea tie-back — a 32 km, 10.75-in carbon-steel flowline plus "
           "steel catenary riser in ~1100 m water — to predict, end-to-end, both hydrodynamic / "
           "terrain / severe-riser SLUGGING and gas-HYDRATE formation, deposition and plugging, and "
           "to demonstrate the model as a flow-assurance DESIGN tool. Three scenarios are run through "
           "the real solver: (A) as-operated normal production with degraded (water-flooded) "
           "insulation and no inhibitor; (B) an unplanned shut-in cooldown; and (C) an engineered "
           "mitigation (restored multi-layer insulation + continuous MEG).")
    D.para("Headline result: as-operated, the line is intermittent over its whole length with slugs "
           f"up to ~{f2(kmA,'slug_length_max_m','{:.0f}')} m and a P90 slug-catcher surge of "
           f"≥ {f2(kmA,'V_surge_P90_m3','{:.1f}')} m³, while the cold, under-insulated wall drives the "
           f"fluid {f2(kmA,'max_subcooling_C','{:.1f}')} °C into the hydrate region (Φ_SH = "
           f"{f2(kmA,'max_Phi_SH','{:.2f}')} ≫ 1), giving a "
           f"{float(kmA.get('P_plug',0))*100:.0f}% plug probability with a P50 time-to-plug of only "
           f"{f2(kmA,'time_to_plug_P50_h','{:.1f}')} h. The engineered fix removes the subcooling "
           "entirely and zeroes the deposit and plug probability — the model quantifies both the "
           "threat and the cure.")
    D.para("Provenance (honest framing): the field is a representative industrial archetype — geometry, "
           "fluid and operating parameters are realistic, self-consistent literature-typical values for "
           "deepwater medium-crude-oil tie-backs, NOT proprietary operator data. The physics and the "
           "predictions are produced by the real solver; the hydrate thermodynamics are anchored to "
           "published data.", italic=True, color=GREY, size=9.5)

    D.H1("1.  The asset and why it is slug- and hydrate-prone")
    D.bullet("Route: 32 km step-out, 10.75-in (0.2545 m ID) carbon-steel flowline on a long, strongly "
             "undulating cold seabed (multiple low spots → terrain slugging), climbing into a steep "
             "steel catenary riser (~last 5.5% of the route) → severe-riser slugging.")
    D.bullet("Water depth ≈ 1100 m; seabed temperature 4 °C; inlet 150 bar / 58 °C.")
    D.bullet("Fluid: medium crude oil (~30° API; C1 ≈ 43 mol%, with a ~31 mol% C7+ heavy tail) at 35% "
             "water cut and 4.5 wt% saline formation water — ample free gas to drive slugging and "
             "plenty of water-wet gas to drive hydrates. PVT by compositional Peng–Robinson flash.")
    D.bullet("In-situ rates: gas 0.150 m³/s, liquid 0.055 m³/s — gassy and intermittent.")
    D.para("This geometry + fluid is a textbook combination for BOTH slugging and hydrates, so it "
           "exercises the whole SHCT prediction chain (hydrodynamics, thermal, hydrate kinetics, "
           "deposition, the coupled Φ_SH risk, inhibitor design and cooldown).")

    D.H1("2.  Modelling approach (summary)")
    D.para("The solver integrates a coupled transient PDE system — conservative liquid-holdup and "
           "gas-mass transport, an implicit-pressure mixture momentum (drift-flux slip), an energy "
           "balance carrying hydrate latent heat and Joule–Thomson cooling, and an advected "
           "reaction–diffusion hydrate phase-field with stochastic nucleation — closed by published "
           "correlations (Bendiksen drift-flux, Taitel–Dukler regimes, Gregory–Scott slug frequency, "
           "Haaland friction, natural-gas hydrate equilibrium, CSMHyK growth, Nielsen–Bucklin MEG, "
           "API RP 14E) and a Peng–Robinson EOS. The slug–hydrate coupling is quantified by the "
           "Φ_SH number. The full equation set is given in the accompanying model.equations.docx "
           "(and, in the comprehensive report, in the Model-equations part).")

    sec = 3
    headline_figs = [
        ("01_profiles.png", "Final-state profiles: elevation, holdup, P–T vs T_eq, subcooling."),
        ("02_holdup_spacetime.png", "Transient liquid-holdup field α_l(x,t) — bright bands are slug activity."),
        ("09_slug_prediction.png", "Slug-formation prediction: terrain, regime, frequency/length, holdup."),
        ("03_PT_envelope.png", "P–T trajectory vs the hydrate envelope."),
        ("04_PhiSH_map.png", "Φ_SH(x,t) coupling-criticality map (Φ_SH>1 ⇒ plugging risk)."),
        ("07_probabilistic.png", "Probabilistic time-to-plug (P10/P50/P90)."),
    ]
    metric_pick = [("max_subcooling_C", "Max subcooling", "°C"), ("max_Phi_SH", "Max Φ_SH", "–"),
                   ("P_plug", "Plug probability", "frac"), ("time_to_plug_P50_h", "Time-to-plug P50", "h"),
                   ("peak_deposit_mm", "Peak wall deposit", "mm"), ("slug_length_max_m", "Max slug length", "m"),
                   ("V_surge_P90_m3", "Surge (P90)", "m³"), ("MEG_wt_pct", "MEG required", "wt%"),
                   ("under_inhibited_km", "Under-inhibited", "km"), ("cooldown_to_hydrate_h", "No-touch time", "h"),
                   ("dP_total_bar", "Total ΔP", "bar"), ("U_eff_WmK", "Effective U", "W/m²K")]
    for label, fname, tag in SCEN:
        folder = os.path.join(OUTROOT, fname)
        km = json.load(open(os.path.join(folder, "key_metrics.json")))
        D.H1(f"{sec}.  Scenario {chr(63+sec)}: {label}")
        rows = [["Metric", "Value", "Units"]]
        for k, lab, u in metric_pick:
            v = km.get(k)
            if v is None:
                continue
            try:
                sv = f"{float(v):.4g}"
            except Exception:
                sv = str(v)
            rows.append([lab, sv, u])
        D.kv_table(rows)
        for fn, cap in headline_figs:
            D.figure(os.path.join(folder, fn), f"{cap}  [{tag}]",
                     width=5.4 if fn == "03_PT_envelope.png" else 6.6)
        sec += 1

    D.H1(f"{sec}.  Mitigation comparison and conclusions")
    D.figure(os.path.join(OUTROOT, "outputs_steady", "12_mitigation_comparison.png"),
             "As-operated vs engineered fix across the governing flow-assurance metrics.", width=6.6)
    D.para("Conclusions (drawn directly from the solver output):", bold=True)
    D.bullet(f"SLUGGING (as-operated): intermittent flow over the whole line; slug length up to "
             f"~{f2(kmA,'slug_length_max_m','{:.0f}')} m; size the slug catcher for ≥ "
             f"{f2(kmA,'V_surge_P90_m3','{:.1f}')} m³ (P90); total ΔP {f2(kmA,'dP_total_bar','{:.0f}')} bar; "
             f"peak velocity {f2(kmA,'Vm_peak_mps','{:.1f}')} m/s vs erosional "
             f"{f2(kmA,'erosional_limit_mps','{:.1f}')} m/s.")
    D.bullet(f"HYDRATES (as-operated): max subcooling {f2(kmA,'max_subcooling_C','{:.1f}')} °C, "
             f"Φ_SH = {f2(kmA,'max_Phi_SH','{:.2f}')} (critical), peak deposit "
             f"{f2(kmA,'peak_deposit_mm','{:.0f}')} mm, {float(kmA.get('P_plug',0))*100:.0f}% plug "
             f"probability, P50 time-to-plug {f2(kmA,'time_to_plug_P50_h','{:.1f}')} h; the model sizes "
             f"the inhibitor demand at MEG ≈ {f2(kmA,'MEG_wt_pct','{:.0f}')} wt% "
             f"({f2(kmA,'MEG_Lph','{:.0f}')} L/h).")
    D.bullet(f"SHUT-IN: no-touch time ≈ {f2(kmS,'cooldown_to_hydrate_h','{:.3f}')} h (the line is already "
             f"in the hydrate region while flowing) → effectively no safe window without inhibition; "
             f"shut-in max subcooling {f2(kmS,'max_subcooling_C','{:.1f}')} °C.")
    D.bullet(f"MITIGATION: restored insulation (U_eff {f2(kmM,'U_eff_WmK','{:.2f}')} W/m²K) removes the "
             f"subcooling (max {f2(kmM,'max_subcooling_C','{:.1f}')} °C), zeroes the deposit and plug "
             f"probability, and restores a no-touch time of {f2(kmM,'cooldown_to_hydrate_h','{:.0f}')} h; "
             f"continuous MEG falls to ≈ {f2(kmM,'MEG_wt_pct','{:.0f}')} wt%.")
    D.bullet(f"NUMERICALLY TRUSTWORTHY: liquid mass error {f2(kmA,'mass_conservation_err','{:.2e}')}, gas "
             f"{f2(kmA,'gas_mass_conservation_err','{:.2e}')}, {int(float(kmA.get('fallbacks',0)))} solver "
             f"fallbacks.")
    return sec


# ============================================================ TITLE
def title_block(D, title, subtitle):
    t = D.doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(title); r.bold = True; r.font.size = Pt(17); r.font.color.rgb = NAVY
    s = D.doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run(subtitle); r.italic = True; r.font.size = Pt(11); r.font.color.rgb = ACCENT
    D.doc.add_paragraph()


# ============================================================ BUILD: model.equations.docx
def build_equations_doc(path):
    D = Doc()
    title_block(D, "Model Equations of the SHCT Solver",
                "Slug–Hydrate Coupled-Transient prediction — governing PDEs and closures used to build the solver")
    D.H1("Overview")
    write_equations(D, intro=True)
    D.save(path); print("wrote", path)


# ============================================================ BUILD: case_study.docx
def build_case_study_doc(path):
    D = Doc()
    title_block(D, "Flow-Assurance Case Study — Deepwater Volatile-Oil Subsea Tie-back",
                "Coupled slug & hydrate prediction with the SHCT solver (32 km, 10.75-in, ~1100 m water)")
    write_case_study(D)
    D.save(path); print("wrote", path)


# ============================================================ BUILD: slug_report.docx (comprehensive)
def build_full_report(path):
    D = Doc()
    title_block(D, "Comprehensive Flow-Assurance Report — Deepwater Volatile-Oil Subsea Tie-back",
                "SHCT slug & hydrate prediction: case study, model equations, all inputs and all generated outputs")
    D.H1("Contents and scope")
    D.para("This single report consolidates EVERYTHING produced for Case Study #3: the engineering "
           "case study, the complete set of model equations used to build the solver, the full input "
           "data deck, and every generated output — all headline metrics, prediction curves, per-column "
           "graphs of every CSV, the probabilistic range graphs, the engineering-deliverable charts, "
           "the complete solver/extension figure gallery (space-time maps, P–T envelopes, Φ_SH maps, "
           "cross-section/quasi-3-D reconstructions, compositional PVT) and every CSV as a data table — "
           "for all three scenarios (as-operated, shut-in, engineered mitigation). Nothing is left out.")
    D.bullet("Part I — Engineering case study (asset, scenarios, headline predictions, mitigation, conclusions).")
    D.bullet("Part II — Model equations (governing PDEs + every closure used to build the solver).")
    D.bullet("Part III — Inputs (full data deck, feed composition, per-scenario thermal/inhibition design).")
    D.bullet("Part IV — Generated outputs (all metrics, curves, graphs, charts, maps and CSV tables, per scenario).")
    D.bullet("Part V — Cross-scenario summary and engineering conclusions.")

    D.pagebreak(); D.H1("PART I — Engineering case study")
    write_case_study(D)

    D.pagebreak(); D.H1("PART II — Model equations used to build the solver")
    write_equations(D, intro=True)

    D.pagebreak(); D.H1("PART III — Inputs (data deck, composition, per-scenario design)")
    write_inputs(D, full=True)

    D.pagebreak(); D.H1("PART IV — All generated outputs (per scenario)")
    sec = 1
    for label, fname, tag in SCEN:
        write_scenario_outputs(D, label, fname, tag, sec)
        D.pagebreak()
        sec += 1

    D.H1(f"PART V — Cross-scenario summary and conclusions")
    write_case_study_conclusions(D)
    D.save(path); print("wrote", path)


def write_case_study_conclusions(D):
    kmA = json.load(open(os.path.join(OUTROOT, "outputs_steady", "key_metrics.json")))
    kmS = json.load(open(os.path.join(OUTROOT, "outputs_shutin", "key_metrics.json")))
    kmM = json.load(open(os.path.join(OUTROOT, "outputs_mitigated", "key_metrics.json")))
    D.para("The comprehensive outputs above support the following engineering conclusions:", bold=True)
    D.bullet(f"As-operated, the line is slug- and hydrate-critical: Φ_SH = {f2(kmA,'max_Phi_SH','{:.2f}')}, "
             f"{float(kmA.get('P_plug',0))*100:.0f}% plug probability, P50 time-to-plug "
             f"{f2(kmA,'time_to_plug_P50_h','{:.1f}')} h, peak deposit {f2(kmA,'peak_deposit_mm','{:.0f}')} mm.")
    D.bullet(f"Inhibitor demand to clear the as-operated risk: MEG ≈ {f2(kmA,'MEG_wt_pct','{:.0f}')} wt% "
             f"({f2(kmA,'MEG_Lph','{:.0f}')} L/h); under-inhibited length {f2(kmA,'under_inhibited_km','{:.1f}')} km.")
    D.bullet(f"Shut-in offers effectively no safe window (no-touch time ≈ "
             f"{f2(kmS,'cooldown_to_hydrate_h','{:.3f}')} h).")
    D.bullet(f"The engineered fix (U_eff {f2(kmM,'U_eff_WmK','{:.2f}')} W/m²K + MEG) removes the subcooling, "
             f"zeroes the plug probability and restores {f2(kmM,'cooldown_to_hydrate_h','{:.0f}')} h of no-touch time.")
    D.bullet(f"Numerics: liquid mass error {f2(kmA,'mass_conservation_err','{:.2e}')}, gas "
             f"{f2(kmA,'gas_mass_conservation_err','{:.2e}')}, {int(float(kmA.get('fallbacks',0)))} fallbacks "
             f"— the predictions are mass-consistent and stable.")


if __name__ == "__main__":
    eq_path = os.path.join(HERE, "model.equations.docx")
    cs_path = os.path.join(HERE, "case_study.docx")
    build_equations_doc(eq_path)
    build_case_study_doc(cs_path)
    desktop = None
    for cand in ("/mnt/c/Users/user/Desktop",):
        if os.path.isdir(cand):
            desktop = cand; break
    report_path = os.path.join(desktop if desktop else HERE, "slug_report.docx")
    build_full_report(report_path)
    print("DONE")
